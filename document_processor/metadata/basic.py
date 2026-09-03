"""基础元数据提取器模块。

本模块实现 :class:`BasicMetadataExtractor`，提供「零额外配置」的文件级元数据
提取能力：对任意文件都能提取文件名、类型、大小、修改时间等通用字段，并对
PDF / Word 等常见格式按需追加页码、章节等结构化字段。设计要点：

- **懒导入重依赖**：``pypdf`` 与 ``python-docx`` 依赖体积较大，且并非所有
  使用场景都需要真正解析文档内容。因此本模块**不在模块顶部 import** 这些
  库，而是分别延迟到对应分支（PDF / docx）的解析代码处执行，避免未安装时
  import 即报错。

- **统一异常契约**：任何提取失败（文件不存在、内容损坏、字段非法等）都会
  统一捕获并转抛为
  :class:`document_processor.exceptions.MetadataExtractionError`，并保留原始
  异常链（``raise ... from exc``），便于调用方定位根因。

- **按文件类型追加字段**：所有文件都具备 ``filename`` / ``file_type`` /
  ``size`` / ``modified_time`` 四类基础字段；在此基础上：

  - PDF（``.pdf``）追加 ``page_count``（页码数，``int``）。
  - Word（``.docx``）追加 ``sections``（章节标题列表，``list[str]``，保序
    去重）。
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from document_processor._utils import _DEFAULT_MAX_FILE_SIZE, is_heading_style
from document_processor.exceptions import MetadataExtractionError
from document_processor.metadata.base import MetadataExtractor


class BasicMetadataExtractor(MetadataExtractor):
    """基础元数据提取器。

    从文件路径出发，提取文件级通用元数据（文件名 / 类型 / 大小 / 修改时间），
    并按文件类型追加结构化字段：PDF 追加页码数，Word 追加章节标题列表。

    适合作为默认 / 兜底的元数据提取实现：无需任何外部配置或模型，仅依赖文件
    系统信息与轻量解析库（``pypdf`` / ``python-docx``）。对于需要更丰富溯源
    信息的场景，可替换为 unstructured 等深度提取后端（后续补充）。

    Example:
        >>> extractor = BasicMetadataExtractor()
        >>> extractor.extract("docs/report.pdf")
        {'filename': 'report.pdf', 'file_type': 'pdf', 'size': 123456,
         'modified_time': '2024-01-01T12:00:00', 'page_count': 10}

        >>> extractor.extract("docs/notes.docx")
        {'filename': 'notes.docx', 'file_type': 'docx', 'size': 23456,
         'modified_time': '2024-01-01T12:00:00',
         'sections': ['第一章', '1.1 背景']}

        >>> extractor.extract("docs/readme.txt")
        {'filename': 'readme.txt', 'file_type': 'txt', 'size': 100,
         'modified_time': '2024-01-01T12:00:00'}
    """

    def __init__(self, max_file_size: int = _DEFAULT_MAX_FILE_SIZE) -> None:
        """初始化基础元数据提取器。

        Args:
            max_file_size: 允许的最大文件字节数，默认 50 MiB（防解压炸弹）。
                提取前会校验实际文件大小，超过该值将抛出
                :class:`MetadataExtractionError`。
        """
        self._max_file_size: int = max_file_size
        """允许的最大文件字节数（防解压炸弹）。"""
    def extract(self, path: str | Path) -> dict[str, Any]:
        """从指定文件提取元数据并返回键值对字典。

        提取流程：

        1. 将 ``path`` 规范化为 :class:`pathlib.Path`，校验其存在且为文件，
           否则抛出 :class:`MetadataExtractionError`。
        2. 校验文件大小不超过 ``self._max_file_size``（防解压炸弹），超限
           抛出 :class:`MetadataExtractionError`。
        3. 提取所有文件都具备的基础字段：``filename``（文件名）、
           ``file_type``（去掉点号的小写后缀）、``size``（字节数）、
           ``modified_time``（ISO 8601 格式的修改时间字符串）。
        4. 按 ``file_type`` 追加类型专属字段：

           - ``"pdf"``：方法体内**懒导入** ``from pypdf import PdfReader``，
             使用 ``PdfReader`` 打开 PDF 并读取 ``len(reader.pages)`` 作为
             ``page_count``。
           - ``"docx"``：方法体内**懒导入** ``import docx``，遍历
             ``doc.Document.paragraphs``，收集样式名以 ``"Heading"`` 开头或
             等于 ``"Title"`` / ``"Subtitle"`` 的段落文本作为 ``sections``，
             结果**保序去重**并跳过空白文本。

        5. 解析过程中任何异常（PDF 损坏、docx 损坏等）都会捕获并转抛为
           :class:`MetadataExtractionError`（``raise ... from exc``），同时
           单独 ``except MetadataExtractionError: raise`` 以避免重复包装。
        6. 返回最终元数据字典。

        Args:
            path: 待提取元数据的文件路径（字符串或 :class:`pathlib.Path`）。

        Returns:
            键值对形式的元数据字典，包含基础字段以及按文件类型追加的
            结构化字段。

        Raises:
            MetadataExtractionError: 文件不存在、不是文件、格式不支持或解析
                失败时抛出。
        """
        p: Path = Path(path)

        if not p.is_file():
            raise MetadataExtractionError(f"目标路径不是有效文件：{p}")

        # 提取前校验文件大小，防止解压炸弹。
        file_size: int = p.stat().st_size
        if file_size > self._max_file_size:
            raise MetadataExtractionError(
                f"文件大小超限：{p}（{file_size} 字节）"
                f"超过允许的最大值 {self._max_file_size} 字节。"
            )

        metadata: dict[str, Any] = {
            "filename": p.name,
            "file_type": p.suffix.lstrip(".").lower(),
            "size": p.stat().st_size,
            "modified_time": datetime.fromtimestamp(p.stat().st_mtime).isoformat(),
        }

        try:
            if metadata["file_type"] == "pdf":
                metadata["page_count"] = self._extract_pdf_page_count(p)
            elif metadata["file_type"] == "docx":
                metadata["sections"] = self._extract_docx_sections(p)
        except MetadataExtractionError:
            raise
        except Exception as exc:
            raise MetadataExtractionError(
                f"提取文件元数据失败：{p}，原因：{exc}"
            ) from exc

        return metadata

    def _extract_pdf_page_count(self, p: Path) -> int:
        """解析 PDF 文件并返回页码数（内部方法）。

        方法体内**懒导入** ``pypdf``（``from pypdf import PdfReader``），使用
        ``PdfReader`` 打开 PDF 并返回 ``len(reader.pages)``。

        Args:
            p: 已规范化且通过存在性校验的 PDF 文件路径。

        Returns:
            PDF 文件的总页数（``int``）。

        Raises:
            MetadataExtractionError: PDF 内容损坏或解析失败时抛出（由调用方
                统一捕获并包装）。
        """
        from pypdf import PdfReader

        reader: PdfReader = PdfReader(str(p))
        return len(reader.pages)

    def _extract_docx_sections(self, p: Path) -> list[str]:
        """解析 Word 文档并返回章节标题列表（内部方法）。

        方法体内**懒导入** ``docx``（``import docx``），遍历
        ``doc.Document.paragraphs``，收集样式名以 ``"Heading"`` 开头或等于
        ``"Title"`` / ``"Subtitle"`` 的段落文本，结果**保序去重**并跳过空白
        文本。

        Args:
            p: 已规范化且通过存在性校验的 Word 文件路径。

        Returns:
            保序去重后的章节标题列表（``list[str]``）。

        Raises:
            MetadataExtractionError: Word 内容损坏或解析失败时抛出（由调用方
                统一捕获并包装）。
        """
        import docx

        doc = docx.Document(str(p))
        sections: list[str] = []
        for paragraph in doc.paragraphs:
            style_name: str | None = (
                paragraph.style.name if paragraph.style is not None else None
            )
            if is_heading_style(style_name):
                text: str = paragraph.text or ""
                # 跳过空白标题文本，避免产生无效章节。
                if not text.strip():
                    continue
                # 保序去重：只追加尚未出现过的标题文本。
                if text not in sections:
                    sections.append(text)
        return sections
