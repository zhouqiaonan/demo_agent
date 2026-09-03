"""Word 文档加载器模块。

本模块基于 ``python-docx`` 库实现 :class:`WordLoader`，用于把 ``.docx`` 文件
按段落解析为 :class:`vector_store.document.Document` 列表。设计要点：

- **懒导入 docx**：``python-docx`` 依赖体积较大，且并非所有使用场景都需要
  真正解析 Word（例如仅引用抽象接口或做单元测试时）。因此本模块**不在模块
  顶部 import docx**，而是把 ``import docx`` 延迟到 :meth:`WordLoader.load`
  方法体内执行，避免未安装 python-docx 时 import 即报错。

- **按段落抽取 + 跳过空白段**：遍历 ``docx.Document.paragraphs``，对每个
  段落取 ``p.text``，遇到文本为空（``None`` 或去除空白后为空）的段落直接
  跳过，避免在下游产生空 chunk。

- **章节追踪**：通过段落样式名（``p.style.name``）识别标题段落（样式名以
  ``"Heading"`` 开头，或等于 ``"Title"`` / ``"Subtitle"``），维护
  ``current_section`` 变量，使后续正文段落的元数据中带上 ``section`` 归属。
  标题段落本身既作为章节标题又作为内容，会作为普通 ``Document`` 产出，
  同时更新后续段落所属章节。

- **0-based 段落索引**：每个 ``Document`` 的元数据中 ``paragraph_index``
  字段使用 ``enumerate(doc.paragraphs)`` 的 **0-based** 索引，与 PDF 加载器
  中 ``page`` 字段的 0-based 语义保持一致，便于下游跨组件对齐定位信息。
"""

from __future__ import annotations

from pathlib import Path

from document_processor._utils import _DEFAULT_MAX_FILE_SIZE, is_heading_style
from document_processor.exceptions import DocumentLoadError
from document_processor.loaders.base import DocumentLoader
from vector_store.document import Document


class WordLoader(DocumentLoader):
    """Word（``.docx``）文件加载器。

    将指定路径的 ``.docx`` 文件按段落解析为一组 :class:`Document`，每个有效
    段落对应一个 ``Document``，其正文为该段落的文本，元数据包含 ``source``
    （来源路径）、``paragraph_index``（0-based 段落索引），并在识别到标题
    段落时额外携带 ``section``（当前段落所属章节标题）。

    Attributes:
        _path: 规范化后的 Word 文件路径（``pathlib.Path``）。
        _max_file_size: 允许的最大文件字节数（默认 50 MiB，防解压炸弹）。

    Example:
        >>> loader = WordLoader("docs/example.docx")
        >>> docs = loader.load()
        >>> docs[0].metadata
        {'source': 'docs/example.docx', 'paragraph_index': 0, 'section': '第一章'}
    """

    def __init__(
        self, path: str | Path, max_file_size: int = _DEFAULT_MAX_FILE_SIZE
    ) -> None:
        """初始化 Word 加载器。

        Args:
            path: Word 文件的路径（字符串或 :class:`pathlib.Path`）。
            max_file_size: 允许的最大文件字节数，默认 50 MiB（防解压炸弹）。
                解析前会校验实际文件大小，超过该值将抛出
                :class:`DocumentLoadError`。
        """
        self._path: Path = Path(path)
        """规范化后的 Word 文件路径。"""

        self._max_file_size: int = max_file_size
        """允许的最大文件字节数（防解压炸弹）。"""

    def load(self) -> list[Document]:
        """解析 Word 文件，按段落抽取文本并返回 :class:`Document` 列表。

        解析流程：

        1. 方法体内**懒导入** ``docx``（``import docx``）。
        2. 校验 ``self._path`` 是否存在且为文件，否则抛出
           :class:`DocumentLoadError`。
        3. 校验文件大小不超过 ``self._max_file_size``（防解压炸弹），超限
           抛出 :class:`DocumentLoadError`。
        4. 使用 ``docx.Document`` 打开 Word 文档，遍历 ``doc.paragraphs``。
        5. 通过段落样式名 ``p.style.name`` 识别标题段落（样式名以
           ``"Heading"`` 开头或等于 ``"Title"`` / ``"Subtitle"``），并更新
           ``current_section``。仅在标题文本去除空白后非空时才更新，避免
           空白标题把 ``current_section`` 污染为空字符串。注意 ``p.style``
           可能为 ``None``，需安全处理。
        6. 跳过文本为空（去除空白后为空）的段落。
        7. 对每个有效段落构造一个 ``Document``，元数据包含 ``source`` 与
           ``paragraph_index``（**0-based** 索引）；若当前已有章节标题
           （``current_section is not None``），则额外追加 ``section`` 字段。
        8. 解析过程中任何异常都会捕获并转抛为 :class:`DocumentLoadError`。

        Returns:
            按段落顺序排列的文档列表；若文档中所有段落均为空白，返回空列表。

        Raises:
            DocumentLoadError: 文件不存在、不是文件，或 Word 解析失败时抛出。
        """
        import docx

        if not self._path.exists():
            raise DocumentLoadError(f"Word 文件不存在：{self._path}")
        if not self._path.is_file():
            raise DocumentLoadError(f"目标路径不是文件：{self._path}")

        # 解析前校验文件大小，防止解压炸弹。
        file_size: int = self._path.stat().st_size
        if file_size > self._max_file_size:
            raise DocumentLoadError(
                f"Word 文件大小超限：{self._path}（{file_size} 字节）"
                f"超过允许的最大值 {self._max_file_size} 字节。"
            )

        documents: list[Document] = []
        # 当前所属章节标题；尚未遇到任何标题段落时为 None。
        current_section: str | None = None
        try:
            doc = docx.Document(str(self._path))
            for paragraph_index, p in enumerate(doc.paragraphs):
                text: str = p.text or ""
                # 识别标题段落：仅在标题文本去除空白后非空时才更新章节归属，
                # 避免空白标题段落把 current_section 污染为空字符串。
                style_name: str | None = p.style.name if p.style is not None else None
                if is_heading_style(style_name) and text.strip():
                    current_section = text
                # 跳过空白段落，避免产生空 chunk。
                if not text.strip():
                    continue
                metadata: dict = {
                    "source": str(self._path),
                    "paragraph_index": paragraph_index,
                }
                if current_section is not None:
                    metadata["section"] = current_section
                documents.append(
                    Document(
                        text=text,
                        metadata=metadata,
                    )
                )
        except DocumentLoadError:
            raise
        except Exception as exc:
            raise DocumentLoadError(
                f"解析 Word 文件失败：{self._path}，原因：{exc}"
            ) from exc

        return documents
